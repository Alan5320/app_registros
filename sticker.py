from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
from tkinter import messagebox
import traceback
from pathlib import Path
import io
import qrcode
from qrcode.constants import ERROR_CORRECT_M


def crear_qr_buffer(texto: str):
    """
    Genera un QR en memoria y devuelve un BytesIO listo para insertarse en Word.
    """
    if not texto:
        raise ValueError("El texto para el QR está vacío")

    qr = qrcode.QRCode(
        version=None,
        error_correction=ERROR_CORRECT_M,
        box_size=12,
        border=1,
    )
    qr.add_data(texto)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _remove_paragraph(paragraph):
    """Elimina un párrafo del documento (uso interno)."""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def _crear_parrafo_centrado(doc):
    """Crea un párrafo centrado y sin espacios extra."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    return p


def _agregar_campo(doc, clave, valor, imprimir_con_titulos, tamaño_fuente_general):
    """Agrega un campo en un párrafo independiente."""
    parrafo = _crear_parrafo_centrado(doc)

    if imprimir_con_titulos:
        run_clave = parrafo.add_run(f"{clave}: ")
        run_clave.bold = True
        run_clave.font.size = tamaño_fuente_general

        run_valor = parrafo.add_run(str(valor))
        run_valor.font.size = tamaño_fuente_general
    else:
        run_valor = parrafo.add_run(str(valor))
        run_valor.font.size = tamaño_fuente_general

    return parrafo


def crear_documento(datos_json, imprimir_con_titulos, imprimir_con_qr, documento_qr):
    try:
        # Validación de datos de entrada
        if not datos_json or not isinstance(datos_json, dict):
            raise ValueError("El JSON de entrada no es válido o está vacío")

        # Copia para no mutar el original
        datos = dict(datos_json)

        # Rutas
        try:
            try:
                ruta_base = Path(__file__).parent
            except NameError:
                ruta_base = Path.cwd()

            ruta_plantilla = (ruta_base / "_internal/plantilla_fija.docx").resolve(strict=True)
            ruta_salida_word = (ruta_base / "sticker_a_imprimir.docx").resolve()
        except FileNotFoundError as e:
            messagebox.showerror("Error", f"No se encontró la plantilla: {e}")
            return False
        except Exception as e:
            messagebox.showerror("Error", f"Error configurando rutas: {e}")
            return False

        # Verificar si el archivo de salida está en uso
        if ruta_salida_word.exists():
            try:
                tmp = ruta_salida_word.with_suffix(".tmp_check")
                ruta_salida_word.rename(tmp)
                tmp.rename(ruta_salida_word)
            except Exception:
                messagebox.showerror(
                    "Archivo en uso",
                    f"El archivo {ruta_salida_word.name} está abierto. Ciérralo y vuelve a intentar."
                )
                return False

        # Cargar plantilla
        doc = Document(ruta_plantilla)

        # Evitar herencia rara: tamaño de fuente 0 a todo
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(0)

        # Márgenes
        for section in doc.sections:
            section.top_margin = Cm(0)
            section.bottom_margin = Cm(0)
            section.left_margin = Cm(0.4)
            section.right_margin = Cm(0.4)

        cantidad_datos = len(datos)

        # Tamaño de página
        seccion = doc.sections[0]
        ancho_pagina = seccion.page_width
        alto_pagina = seccion.page_height

        # Valores por defecto por si la plantilla no cae en ninguna condición
        tamaño_fuente_nombres = Pt(12)
        tamaño_fuente_general = Pt(10)

        # ----------------------------- Tamaño (8X6) -----------------------------
        if ancho_pagina > Cm(8) and alto_pagina > Cm(6):

            if imprimir_con_qr:
                margen_superior = {
                    1: Cm(0.8), 2: Cm(0.7), 3: Cm(0.4), 4: Cm(0.3),
                }.get(cantidad_datos, Cm(0.3))

                ajustes_fuente = {
                    1: (Pt(22), Pt(22)),
                    2: (Pt(18), Pt(18)),
                    3: (Pt(18), Pt(16)),
                    4: (Pt(17), Pt(14)),
                    5: (Pt(16), Pt(13)),
            }
            else:
                margen_superior = {
                    1: Cm(2), 2: Cm(1.8), 3: Cm(1.6), 4: Cm(1.4),
                }.get(cantidad_datos, Cm(1.1))
                ajustes_fuente = {
                    1: (Pt(22), Pt(22)),
                    2: (Pt(20), Pt(18)),
                    3: (Pt(18), Pt(15)),
                    4: (Pt(17), Pt(14)),
                    5: (Pt(16), Pt(13)),
                }

            for section in doc.sections:
                section.top_margin = margen_superior
                section.bottom_margin = Cm(0.2)


            if cantidad_datos >= 6:
                tamaño_fuente_nombres, tamaño_fuente_general = (Pt(15), Pt(11))
            else:
                tamaño_fuente_nombres, tamaño_fuente_general = ajustes_fuente[cantidad_datos]

        # ----------------------------- Tamaño (10x5) -----------------------------
        elif ancho_pagina > Cm(10) and alto_pagina > Cm(5):

            if imprimir_con_qr:
                margen_superior = {
                    1: Cm(0.7), 2: Cm(0.5), 3: Cm(0.3), 4: Cm(0.3),
                }.get(cantidad_datos, Cm(0.3))

                ajustes_fuente = {
                    1: (Pt(22), Pt(22)),
                    2: (Pt(18), Pt(18)),
                    3: (Pt(16), Pt(15)),
                    4: (Pt(15), Pt(14)),
                    5: (Pt(14), Pt(13)),
            }
            else:
                margen_superior = {
                    1: Cm(1.6), 2: Cm(1.4), 3: Cm(1.2), 4: Cm(1),
                }.get(cantidad_datos, Cm(1.1))
                ajustes_fuente = {
                    1: (Pt(22), Pt(22)),
                    2: (Pt(20), Pt(18)),
                    3: (Pt(18), Pt(15)),
                    4: (Pt(17), Pt(14)),
                    5: (Pt(16), Pt(13)),
                }

            for section in doc.sections:
                section.top_margin = margen_superior
                section.bottom_margin = Cm(0.2)

            if cantidad_datos >= 6:
                tamaño_fuente_nombres, tamaño_fuente_general = (Pt(15), Pt(11))
            else:
                tamaño_fuente_nombres, tamaño_fuente_general = ajustes_fuente[cantidad_datos]

        # ----------------------------- Tamaño (6x3) -----------------------------
        elif ancho_pagina > Cm(6) and alto_pagina > Cm(3):
            margen_superior = {
                1: Cm(0.8), 2: Cm(0.7), 3: Cm(0.6), 4: Cm(0.5),
            }.get(cantidad_datos, Cm(0.4))

            for section in doc.sections:
                section.top_margin = margen_superior
                section.bottom_margin = Cm(0.2)

            ajustes_fuente = {
                1: (Pt(14), Pt(12)),
                2: (Pt(13), Pt(11)),
                3: (Pt(12), Pt(10)),
                4: (Pt(11), Pt(9)),
                5: (Pt(10), Pt(8)),
            }

            if cantidad_datos >= 6:
                tamaño_fuente_nombres, tamaño_fuente_general = (Pt(9), Pt(7))
            else:
                tamaño_fuente_nombres, tamaño_fuente_general = ajustes_fuente[cantidad_datos]

        # Limpiar párrafos vacíos de la plantilla
        for p in list(doc.paragraphs):
            if not p.text.strip():
                _remove_paragraph(p)

        # NOMBRES primero
        nombres = datos.pop("NOMBRES", None)
        if nombres:
            parrafo_nombres = _crear_parrafo_centrado(doc)
            run_nombres = parrafo_nombres.add_run(str(nombres))
            run_nombres.font.size = tamaño_fuente_nombres
            run_nombres.font.bold = True

        # APELLIDOS y DOCUMENTO primero
        if "APELLIDOS" in datos:
            valor = datos.pop("APELLIDOS")
            _agregar_campo(doc, "APELLIDOS", valor, imprimir_con_titulos, tamaño_fuente_general)

        if "DOCUMENTO" in datos:
            valor = datos.pop("DOCUMENTO")
            _agregar_campo(doc, "DOCUMENTO", valor, imprimir_con_titulos, tamaño_fuente_general)

        # Resto de campos
        for clave, valor in datos.items():
            _agregar_campo(doc, clave, valor, imprimir_con_titulos, tamaño_fuente_general)

        # QR justo debajo de DOCUMENTO
        if imprimir_con_qr and ancho_pagina > Cm(8) and alto_pagina > Cm(6):
            qr_buffer = crear_qr_buffer(str(documento_qr))

            parrafo_qr = _crear_parrafo_centrado(doc)
            run_qr = parrafo_qr.add_run()
            run_qr.add_picture(qr_buffer, width=Cm(2.5))


        # QR justo debajo de DOCUMENTO
        if imprimir_con_qr and ancho_pagina > Cm(10) and alto_pagina > Cm(5):
            qr_buffer = crear_qr_buffer(str(documento_qr))

            parrafo_qr = _crear_parrafo_centrado(doc)
            run_qr = parrafo_qr.add_run()
            run_qr.add_picture(qr_buffer, width=Cm(2.3))


        # Limpieza final: no borrar párrafos que contengan imagen
        for p in list(doc.paragraphs):
            if (
                not p.text.strip()
                and not p._element.xpath('.//w:drawing')
                and len(doc.paragraphs) > 1
            ):
                _remove_paragraph(p)

        # Guardar documento
        try:
            doc.save(ruta_salida_word)
        except PermissionError:
            messagebox.showerror(
                "Error de permisos",
                f"No se pudo guardar el documento. Asegúrate de tener permisos de escritura en:\n{ruta_salida_word}"
            )
            return False

        return True

    except Exception as e:
        traceback.print_exc()
        messagebox.showerror(
            "Error crítico",
            f"Error inesperado: {str(e)}\nDetalles en la consola."
        )
        return False